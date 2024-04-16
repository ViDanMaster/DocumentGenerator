import os
from docx import Document
from docx.shared import Pt
from Factories.AbstractFactory import AbstractFactory

class WordFactory(AbstractFactory):
    def __init__(self, name):
        self.folder_path = 'Results/'
        self.filename = name
        self.document = Document()

    def createTitle(self, title_text, font_size, font_family, alignment) -> object:
        paragraph = self.document.add_heading(0)
        paragraph.alignment = alignment
        run = paragraph.add_run(title_text)
        run.font.size = Pt(font_size)
        run.font.name = font_family
        return run

    def createContent(self, content_text, font_size, font_family, alignment) -> object:
        paragraph = self.document.add_paragraph()
        paragraph.alignment = alignment
        run = paragraph.add_run(content_text)
        run.font.size = Pt(font_size)
        run.font.name = font_family
        return run

    def createFooter(self, footer_text, font_size, font_family, alignment) -> object:
        footer = self.document.sections[0].footer
        paragraph = footer.add_paragraph()
        paragraph.alignment = alignment
        run = paragraph.add_run(footer_text)
        run.font.size = Pt(font_size)
        run.font.name = font_family
        return run
    
    def save(self):
        self.document.save(self.getSavePath())

    def getSavePath(self):
        return os.path.join(self.folder_path, self.filename)